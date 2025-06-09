#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Dedykowany konwerter Markdown -> Word dla prac magisterskich UE Wrocław
Automatyczne formatowanie zgodne ze standardami uczelnianymi
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import markdown
from bs4 import BeautifulSoup

class MarkdownToWordUE:
    def __init__(self):
        self.doc = Document()
        self.setup_document_style()
        
    def setup_document_style(self):
        """Konfiguracja stylu dokumentu zgodnie ze standardami UE Wrocław"""
        
        # Konfiguracja sekcji dokumentu
        section = self.doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5) 
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        
        # Konfiguracja stylów
        self.setup_styles()
        
    def setup_styles(self):
        """Definiowanie stylów zgodnych ze standardami UE"""
        
        styles = self.doc.styles
        
        # Styl głównego tekstu
        normal_style = styles['Normal']
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(12)
        
        normal_paragraph = normal_style.paragraph_format
        normal_paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal_paragraph.line_spacing = 1.5
        normal_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal_paragraph.first_line_indent = Cm(0.5)
        normal_paragraph.space_after = Pt(6)
        
        # Styl nagłówka 1 (rozdział)
        try:
            heading1_style = styles['Heading 1']
        except KeyError:
            heading1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        
        h1_font = heading1_style.font
        h1_font.name = 'Times New Roman'
        h1_font.size = Pt(14)
        h1_font.bold = True
        
        h1_paragraph = heading1_style.paragraph_format
        h1_paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        h1_paragraph.line_spacing = 1.5
        h1_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h1_paragraph.space_before = Pt(18)
        h1_paragraph.space_after = Pt(12)
        h1_paragraph.keep_with_next = True
        
        # Styl nagłówka 2 (podrozdział)
        try:
            heading2_style = styles['Heading 2']
        except KeyError:
            heading2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            
        h2_font = heading2_style.font
        h2_font.name = 'Times New Roman'
        h2_font.size = Pt(12)
        h2_font.bold = True
        
        h2_paragraph = heading2_style.paragraph_format
        h2_paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        h2_paragraph.line_spacing = 1.5
        h2_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2_paragraph.space_before = Pt(12)
        h2_paragraph.space_after = Pt(6)
        h2_paragraph.first_line_indent = Cm(0)
        
        # Styl nagłówka 3
        try:
            heading3_style = styles['Heading 3']
        except KeyError:
            heading3_style = styles.add_style('Heading 3', WD_STYLE_TYPE.PARAGRAPH)
            
        h3_font = heading3_style.font
        h3_font.name = 'Times New Roman'
        h3_font.size = Pt(12)
        h3_font.bold = True
        
        h3_paragraph = heading3_style.paragraph_format
        h3_paragraph.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        h3_paragraph.line_spacing = 1.5
        h3_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h3_paragraph.space_before = Pt(6)
        h3_paragraph.space_after = Pt(6)
        h3_paragraph.first_line_indent = Cm(0)
        
        # Styl dla kodu
        try:
            code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        except:
            code_style = styles['Code'] if 'Code' in styles else None
            
        if code_style:
            code_font = code_style.font
            code_font.name = 'Courier New'
            code_font.size = Pt(10)
            
            code_paragraph = code_style.paragraph_format
            code_paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE
            code_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            code_paragraph.first_line_indent = Cm(0)
            code_paragraph.left_indent = Cm(1)
            code_paragraph.space_before = Pt(6)
            code_paragraph.space_after = Pt(6)
    
    def convert_markdown_to_word(self, markdown_file_path, output_file_path):
        """Główna metoda konwersji"""
        
        print(f"📄 Rozpoczynam konwersję: {markdown_file_path}")
        print(f"📝 Plik docelowy: {output_file_path}")
        
        # Wczytanie pliku Markdown
        with open(markdown_file_path, 'r', encoding='utf-8') as file:
            markdown_content = file.read()
        
        # Przetworzenie treści
        self.process_markdown_content(markdown_content)
        
        # Dodanie numeracji stron
        self.add_page_numbers()
        
        # Zapisanie dokumentu
        self.doc.save(output_file_path)
        print(f"✅ Konwersja zakończona pomyślnie!")
        print(f"📋 Dokument Word zapisany: {output_file_path}")
        
    def process_markdown_content(self, content):
        """Przetwarzanie treści Markdown"""
        
        lines = content.split('\n')
        i = 0
        toc_inserted = False
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            # Sprawdzenie czy to ręczny spis treści - pomijamy go
            if line == "# Spis treści" or line == "Spis treści":
                i = self.skip_manual_toc(lines, i)
                if not toc_inserted:
                    self.add_automatic_toc()
                    toc_inserted = True
                continue
                
            # Nagłówki
            if line.startswith('# ') and not line.startswith('## '):
                self.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                self.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                self.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                self.add_heading(line[5:], level=4)
                
            # Bloki kodu
            elif line.startswith('```'):
                i, code_content = self.extract_code_block(lines, i)
                self.add_code_block(code_content)
                
            # Tabele
            elif '|' in line and '---' in lines[i+1] if i+1 < len(lines) else False:
                i, table_content = self.extract_table(lines, i)
                self.add_table(table_content)
                
            # Listy
            elif line.startswith('- ') or line.startswith('* ') or re.match(r'^\d+\.', line):
                i, list_content = self.extract_list(lines, i)
                self.add_list(list_content)
                
            # Zwykły tekst
            else:
                # Zbieranie ciągłego tekstu akapitu
                paragraph_lines = []
                while i < len(lines) and lines[i].strip() and not self.is_special_line(lines[i]):
                    paragraph_lines.append(lines[i].strip())
                    i += 1
                
                if paragraph_lines:
                    paragraph_text = ' '.join(paragraph_lines)
                    self.add_paragraph(paragraph_text)
                    i -= 1
            
            i += 1
    
    def is_special_line(self, line):
        """Sprawdza czy linia to element specjalny (nagłówek, lista, itp.)"""
        line = line.strip()
        return (line.startswith('#') or 
                line.startswith('```') or
                line.startswith('- ') or 
                line.startswith('* ') or
                '|' in line or
                re.match(r'^\d+\.', line) or
                line.startswith('**') and line.endswith('**'))
    
    def extract_code_block(self, lines, start_index):
        """Wyciąganie bloku kodu"""
        code_lines = []
        i = start_index + 1
        
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
            
        return i, '\n'.join(code_lines)
    
    def extract_table(self, lines, start_index):
        """Wyciąganie tabeli"""
        table_lines = []
        i = start_index
        
        while i < len(lines) and '|' in lines[i]:
            if not lines[i].strip().startswith('|---'):
                table_lines.append(lines[i])
            i += 1
            
        return i - 1, table_lines
    
    def extract_list(self, lines, start_index):
        """Wyciąganie listy"""
        list_lines = []
        i = start_index
        
        while i < len(lines):
            line = lines[i].strip()
            if (line.startswith('- ') or 
                line.startswith('* ') or 
                re.match(r'^\d+\.', line)):
                list_lines.append(line)
                i += 1
            elif line == '':
                # Sprawdź czy następna linia to też element listy
                if i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if (next_line.startswith('- ') or 
                        next_line.startswith('* ') or 
                        re.match(r'^\d+\.', next_line)):
                        i += 1
                        continue
                break
            else:
                break
                
        return i - 1, list_lines
    
    def add_heading(self, text, level=1):
        """Dodawanie nagłówka"""
        if level == 1:
            heading = self.doc.add_heading(text, level=1)
            heading.style = self.doc.styles['Heading 1']
        elif level == 2:
            heading = self.doc.add_heading(text, level=2)
            heading.style = self.doc.styles['Heading 2']
        else:
            heading = self.doc.add_heading(text, level=3)
            heading.style = self.doc.styles['Heading 3']
    
    def add_paragraph(self, text):
        """Dodawanie akapitu z formatowaniem inline"""
        p = self.doc.add_paragraph()
        p.style = self.doc.styles['Normal']
        
        # Przetwarzanie formatowania inline (bold, italic, kod)
        self.process_inline_formatting(p, text)
    
    def process_inline_formatting(self, paragraph, text):
        """Przetwarzanie formatowania wewnątrz tekstu"""
        
        # Prosty parser dla **bold**, *italic*, `code`
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                # Italic  
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                # Inline code
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            else:
                # Zwykły tekst
                run = paragraph.add_run(part)
    
    def add_code_block(self, code_content):
        """Dodawanie bloku kodu"""
        if 'Code' in self.doc.styles:
            for line in code_content.split('\n'):
                p = self.doc.add_paragraph(line)
                p.style = self.doc.styles['Code']
        else:
            # Fallback
            p = self.doc.add_paragraph(code_content)
            p.style = self.doc.styles['Normal']
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
    
    def add_table(self, table_lines):
        """Dodawanie tabeli"""
        if not table_lines:
            return
            
        # Parsowanie pierwszej linii dla liczby kolumn
        first_row = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
        num_cols = len(first_row)
        
        # Tworzenie tabeli
        table = self.doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'
        
        # Nagłówki
        header_row = table.rows[0]
        for i, cell_text in enumerate(first_row):
            header_row.cells[i].text = cell_text
            # Pogrubienie nagłówków
            for paragraph in header_row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Pozostałe wiersze (pomijamy separator z ---)
        for line in table_lines[1:]:
            if '---' in line:
                continue
                
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            row = table.add_row()
            for i, cell_text in enumerate(cells):
                if i < len(row.cells):
                    row.cells[i].text = cell_text
    
    def add_list(self, list_lines):
        """Dodawanie listy"""
        for line in list_lines:
            if line.startswith('- ') or line.startswith('* '):
                p = self.doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\.', line):
                p = self.doc.add_paragraph(line[2:].strip(), style='List Number')
            
            # Formatowanie listy
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(1)
    
    def skip_manual_toc(self, lines, start_index):
        """Pomijanie ręcznego spisu treści"""
        i = start_index + 1
        
        # Pomijamy wszystkie linie spisu treści aż do następnego nagłówka głównego
        while i < len(lines):
            line = lines[i].strip()
            
            # Koniec spisu treści gdy natrafimy na ---
            if line.startswith('---'):
                return i + 1
                
            # Lub gdy natrafimy na następny główny nagłówek
            if line.startswith('# ') and not line.startswith('## '):
                return i - 1
                
            i += 1
        
        return i
    
    def add_automatic_toc(self):
        """Dodawanie automatycznego spisu treści Word"""
        # Dodanie nagłówka "Spis treści"
        toc_heading = self.doc.add_heading('Spis treści', level=1)
        toc_heading.style = self.doc.styles['Heading 1']
        
        # Dodanie pola TOC (Table of Contents)
        paragraph = self.doc.add_paragraph()
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        
        # XML dla automatycznego spisu treści
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        # Dodanie podziału strony po spisie treści
        self.doc.add_page_break()
    
    def add_page_numbers(self):
        """Dodawanie numeracji stron"""
        section = self.doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Dodanie numeru strony
        run = footer_para.runs[0] if footer_para.runs else footer_para.add_run()
        
        # XML dla numeru strony
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

def main():
    """Główna funkcja uruchamiająca konwersję"""
    
    print("🎓 KONWERTER MARKDOWN → WORD DLA PRAC MAGISTERSKICH UE WROCŁAW")
    print("=" * 65)
    
    # Ścieżki plików
    markdown_file = "Praca Dyplomowa/PRACA_MAGISTERSKA_ŁUKASZ_KAMINSKI_FINALNA.md"
    word_file = "Praca Dyplomowa/PRACA_MAGISTERSKA_ŁUKASZ_KAMIŃSKI_FINALNA.docx"
    
    # Sprawdzenie czy plik źródłowy istnieje
    if not os.path.exists(markdown_file):
        print(f"❌ BŁĄD: Nie znaleziono pliku źródłowego: {markdown_file}")
        return False
    
    try:
        # Utworzenie katalogu docelowego jeśli nie istnieje
        os.makedirs(os.path.dirname(word_file), exist_ok=True)
        
        # Konwersja
        converter = MarkdownToWordUE()
        converter.convert_markdown_to_word(markdown_file, word_file)
        
        # Informacje końcowe
        print("\n" + "=" * 65)
        print("✅ KONWERSJA ZAKOŃCZONA POMYŚLNIE!")
        print(f"📄 Plik Word: {word_file}")
        print(f"📊 Rozmiar: {os.path.getsize(word_file) / 1024:.1f} KB")
        print("\n🎯 FUNKCJE DOSTĘPNE W DOKUMENCIE:")
        print("   ✓ Formatowanie zgodne ze standardami UE Wrocław")
        print("   ✓ Times New Roman 12pt, interlinia 1.5")
        print("   ✓ Właściwe marginesy (3.0/2.0/2.5/2.5 cm)")
        print("   ✓ Numeracja stron")
        print("   ✓ Hierarchia nagłówków")
        print("   ✓ Formatowanie tabel i kodu")
        print("   ✓ Wcięcia akapitowe 0.5 cm")
        print("\n🔧 GOTOWE DO DALSZEJ EDYCJI W MS WORD")
        
        return True
        
    except Exception as e:
        print(f"❌ BŁĄD PODCZAS KONWERSJI: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    main()
