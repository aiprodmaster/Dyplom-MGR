#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Konwerter Markdown do Word dla Pracy Magisterskiej
Łukasz Kamiński - Inteligentne wsparcie procesu wdrożenia systemu ERP przy wykorzystaniu AI
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

class MasterThesisConverter:
    def __init__(self):
        self.doc = Document()
        self.setup_document_styles()
        
    def setup_document_styles(self):
        """Ustawienia zgodne ze standardami UE Wrocław"""
        
        # Ustawienia sekcji i marginesów
        section = self.doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        
        # Style nagłówków
        styles = self.doc.styles
        
        # Styl dla tytułu głównego
        if 'Title Main' not in [s.name for s in styles]:
            title_style = styles.add_style('Title Main', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Times New Roman'
            title_font.size = Pt(16)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # Styl dla nagłówków 1
        heading1 = styles['Heading 1']
        heading1.font.name = 'Times New Roman'
        heading1.font.size = Pt(14)
        heading1.font.bold = True
        heading1.paragraph_format.space_before = Pt(18)
        heading1.paragraph_format.space_after = Pt(6)
        
        # Styl dla nagłówków 2
        heading2 = styles['Heading 2']
        heading2.font.name = 'Times New Roman'
        heading2.font.size = Pt(13)
        heading2.font.bold = True
        heading2.paragraph_format.space_before = Pt(12)
        heading2.paragraph_format.space_after = Pt(6)
        
        # Styl dla nagłówków 3
        heading3 = styles['Heading 3']
        heading3.font.name = 'Times New Roman'
        heading3.font.size = Pt(12)
        heading3.font.bold = True
        heading3.paragraph_format.space_before = Pt(6)
        heading3.paragraph_format.space_after = Pt(3)
        
        # Styl dla tekstu normalnego
        normal = styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(12)
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.first_line_indent = Cm(0.5)
        normal.paragraph_format.line_spacing = 1.5
        normal.paragraph_format.space_after = Pt(6)
        
    def add_title_page(self):
        """Dodaje stronę tytułową zgodną ze standardami UE"""
        
        # Kierunek studiów
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('KIERUNEK STUDIÓW')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('MASTER OF BUSINESS ADMINISTRATION')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        
        # Dodaj odstęp
        self.doc.add_paragraph('\n' * 3)
        
        # Autor
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Łukasz Kamiński\nNr albumu 190139')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        # Dodaj odstęp
        self.doc.add_paragraph('\n' * 2)
        
        # Typ pracy
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('PRACA MAGISTERSKA')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        
        # Dodaj odstęp
        self.doc.add_paragraph('\n' * 2)
        
        # Tytuł pracy
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('INTELIGENTNE WSPARCIE PROCESU WDROŻENIA\nSYSTEMU KLASY ERP PRZY WYKORZYSTANIU AI')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        
        # Dodaj odstęp
        self.doc.add_paragraph('\n' * 4)
        
        # Promotor
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Promotor:\ndr hab. Natalia Szozda, prof. UEW\nKatedra Logistyki')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        # Dodaj odstęp
        self.doc.add_paragraph('\n' * 3)
        
        # Miejsce i rok
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('WROCŁAW 2025')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        
        # Nowa strona
        self.doc.add_page_break()
        
    def add_abstract_page(self):
        """Dodaje stronę z tytułem angielskim i streszczeniem"""
        
        # Tytuł angielski
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('INTELLIGENT SUPPORT FOR ERP SYSTEM\nIMPLEMENTATION PROCESS USING AI')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        
        self.doc.add_paragraph()
        
        # Streszczenie
        heading = self.doc.add_heading('Abstract', level=2)
        heading.style = self.doc.styles['Heading 2']
        
        abstract_text = """This master's thesis investigates the application of artificial intelligence technologies, particularly Retrieval-Augmented Generation (RAG) and natural language processing (NLP), to enhance the efficiency of ERP system implementation processes. The research focuses on developing and testing an AI assistant prototype designed to support ERP consultants during system deployment projects. Through comprehensive literature analysis and practical prototype development, the study demonstrates that AI-powered solutions can significantly improve operational efficiency, reduce implementation time, and enhance decision-making quality in ERP environments. The prototype, developed specifically for Comarch ERP XL integration, showcases the practical applicability of advanced AI technologies in business process optimization. Key findings indicate substantial potential for automation of routine tasks, intelligent data analysis, and personalized user interfaces in ERP implementation scenarios."""
        
        p = self.doc.add_paragraph(abstract_text)
        p.style = self.doc.styles['Normal']
        
        # Słowa kluczowe
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        run = p.add_run('Keywords: ')
        run.font.bold = True
        p.add_run('artificial intelligence, ERP systems, implementation, RAG, automation, digital transformation')
        
        # Nowa strona
        self.doc.add_page_break()
        
    def process_markdown_content(self, content):
        """Przetwarza zawartość markdown i konwertuje na Word"""
        
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Pomiń pustę linie na początku
            if not line:
                i += 1
                continue
            
            # Pomiń stronę tytułową i abstract - już dodane
            if 'KIERUNEK STUDIÓW' in line or 'Abstract' in line:
                while i < len(lines) and not lines[i].startswith('# Spis treści'):
                    i += 1
                continue
                
            # Spis treści
            if line.startswith('# Spis treści') or line == 'Spis treści':
                self.add_table_of_contents()
                i += 1
                continue
                
            # Nagłówki
            if line.startswith('#'):
                self.add_heading(line)
                i += 1
                continue
                
            # Tabele
            if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
                table_lines = []
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                if table_lines:
                    self.add_table(table_lines)
                continue
                
            # Listy numerowane i punktorowane
            if re.match(r'^\d+\.', line) or line.startswith('- ') or line.startswith('* '):
                list_items = []
                while i < len(lines) and (re.match(r'^\d+\.', lines[i].strip()) or 
                                         lines[i].strip().startswith('- ') or 
                                         lines[i].strip().startswith('* ')):
                    list_items.append(lines[i].strip())
                    i += 1
                self.add_list(list_items)
                continue
                
            # Zwykły tekst
            if line:
                paragraph_text = line
                i += 1
                
                # Zbierz kolejne linie należące do tego samego akapitu
                while i < len(lines) and lines[i].strip() and not self.is_special_line(lines[i]):
                    paragraph_text += ' ' + lines[i].strip()
                    i += 1
                    
                self.add_paragraph(paragraph_text)
            else:
                i += 1
                
    def is_special_line(self, line):
        """Sprawdza czy linia jest specjalna (nagłówek, lista, tabela)"""
        line = line.strip()
        return (line.startswith('#') or 
                line.startswith('- ') or 
                line.startswith('* ') or 
                re.match(r'^\d+\.', line) or
                '|' in line)
                
    def add_heading(self, heading_line):
        """Dodaje nagłówek"""
        level = heading_line.count('#')
        text = heading_line.lstrip('#').strip()
        
        if level == 1:
            heading = self.doc.add_heading(text, level=1)
        elif level == 2:
            heading = self.doc.add_heading(text, level=2)
        else:
            heading = self.doc.add_heading(text, level=3)
            
    def add_paragraph(self, text):
        """Dodaje akapit z formatowaniem"""
        # Usuń markdown formatting
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # bold
        text = re.sub(r'\*(.*?)\*', r'\1', text)      # italic
        text = re.sub(r'`(.*?)`', r'\1', text)        # code
        
        if text.strip():
            p = self.doc.add_paragraph(text)
            p.style = self.doc.styles['Normal']
            
    def add_list(self, list_items):
        """Dodaje listę numerowaną lub punktowaną"""
        for item in list_items:
            text = re.sub(r'^\d+\.\s*', '', item)  # usuń numerację
            text = re.sub(r'^[-*]\s*', '', text)   # usuń punktor
            p = self.doc.add_paragraph(text, style='List Bullet')
            
    def add_table(self, table_lines):
        """Dodaje tabelę"""
        if len(table_lines) < 2:
            return
            
        # Parse table
        rows = []
        for line in table_lines:
            if '---' in line:  # separator line
                continue
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)
                
        if rows:
            table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Table Grid'
            
            for i, row_data in enumerate(rows):
                for j, cell_data in enumerate(row_data):
                    if j < len(table.rows[i].cells):
                        table.rows[i].cells[j].text = cell_data
                        
    def add_table_of_contents(self):
        """Dodaje spis treści"""
        heading = self.doc.add_heading('Spis treści', level=1)
        
        toc_items = [
            ("1. Wstęp", 1),
            ("2. Systemy ERP w zarządzaniu przedsiębiorstwem", 5),
            ("   2.1 Definicja i charakterystyka systemów ERP", 5),
            ("   2.2 Komponenty i architektura systemów ERP", 8),
            ("   2.3 Wyzwania wdrożeniowe systemów ERP", 12),
            ("3. Sztuczna inteligencja jako wsparcie systemów ERP", 18),
            ("   3.1 Fundamenty AI w zastosowaniach biznesowych", 18),
            ("   3.2 Technologia Retrieval-Augmented Generation (RAG)", 24),
            ("   3.3 Zastosowania AI w systemach ERP", 30),
            ("4. Metodyka badań", 36),
            ("   4.1 Problem badawczy i hipotezy", 36),
            ("   4.2 Metodologia projektowania prototypu", 38),
            ("   4.3 Narzędzia i technologie wykorzystane w badaniu", 42),
            ("5. Projektowanie i implementacja prototypu asystenta AI", 46),
            ("   5.1 Specyfikacja funkcjonalna systemu", 46),
            ("   5.2 Architektura rozwiązania", 52),
            ("   5.3 Implementacja kluczowych modułów", 58),
            ("   5.4 Integracja z systemem Comarch ERP XL", 65),
            ("6. Badanie efektywności prototypu", 72),
            ("   6.1 Metodyka testowania", 72),
            ("   6.2 Scenariusze testowe", 75),
            ("   6.3 Analiza wyników", 78),
            ("   6.4 Ocena jakości odpowiedzi systemu RAG", 84),
            ("7. Zakończenie", 88),
            ("Bibliografia", 93),
            ("Wykaz tabel", 99),
            ("Wykaz rysunków", 100),
            ("Załączniki", 101)
        ]
        
        for item, page in toc_items:
            p = self.doc.add_paragraph()
            p.add_run(item)
            
            # Dodaj kropki prowadzące
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Cm(15))
            p.add_run('\t')
            p.add_run(str(page))
            
        self.doc.add_page_break()
        
    def convert_file(self, input_file, output_file):
        """Konwertuje plik markdown na Word"""
        
        print(f"Czytanie pliku: {input_file}")
        with open(input_file, 'r', encoding='utf-8') as f:
            content = f.read()
            
        print("Generowanie strony tytułowej...")
        self.add_title_page()
        
        print("Generowanie strony z abstraktem...")
        self.add_abstract_page()
        
        print("Przetwarzanie zawartości...")
        self.process_markdown_content(content)
        
        print(f"Zapisywanie do pliku: {output_file}")
        self.doc.save(output_file)
        print("Konwersja zakończona pomyślnie!")

def main():
    """Główna funkcja konwertera"""
    
    input_file = "Praca Dyplomowa/PRACA_MAGISTERSKA_ŁUKASZ_KAMINSKI_FINALNA_Z_PROTOTYPEM_KOMPLETNA.md"
    output_file = "Praca Dyplomowa/PRACA_MAGISTERSKA_ŁUKASZ_KAMINSKI_KOMPLETNA.docx"
    
    if not os.path.exists(input_file):
        print(f"BŁĄD: Nie znaleziono pliku wejściowego: {input_file}")
        return
        
    converter = MasterThesisConverter()
    converter.convert_file(input_file, output_file)

if __name__ == "__main__":
    main()
